import os
import platform

from docx import Document
from typing import Any, Collection


__all__ = [
    'get_desktop_path',
    'TableReportGenerator'
]

from docx.shared import Inches


def get_desktop_path() -> str:
    system = platform.system()
    if system == 'Windows':
        return os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
    elif system == 'Darwin' or system == 'Linux':
        return os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop')
    else:
        return ''


class TableReportGenerator:
    def __init__(self, document_name: str, rows_count: int, columns_count: int, headers: list[str],
                 top_margin: int = 0.5, bottom_margin: int = 0.5, left_margin: int = 0.3, right_margin: int = 0.3):
        assert len(headers) == columns_count, 'Количество заголовков и столбцов должны совпадать'

        self.document = Document()
        self.rows_count = rows_count
        self.columns_count = columns_count
        self.headers = headers
        self.document_name = document_name
        self.top_margin = top_margin
        self.bottom_margin = bottom_margin
        self.left_margin = left_margin
        self.right_margin = right_margin

    def add_header(self, name: str, level: int):
        self.document.add_heading(name, level=level)

    def add_paragraph(self, text: str):
        self.document.add_paragraph(text)

    def fill_table(self, record_set: Collection[Collection[Any]]):
        table = self.document.add_table(rows=1, cols=self.columns_count)
        table.style = 'Table Grid'
        table.autofit = True

        cells_headers = table.rows[0].cells
        headers_counter: int = 0

        for header in self.headers:
            cells_headers[headers_counter].text = header
            headers_counter += 1

        for record in record_set:
            row_cells = table.add_row().cells

            for i in range(len(record)):
                row_cells[i].text = str(record[i])

    def save(self, file_path):
        for section in self.document.sections:
            section.top_margin = Inches(self.top_margin)
            section.bottom_margin = Inches(self.bottom_margin)
            section.left_margin = Inches(self.left_margin)
            section.right_margin = Inches(self.right_margin)

        try:
            self.document.save(file_path)
        except FileNotFoundError:
            self.document.save(self.document_name)

